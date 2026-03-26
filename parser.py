"""
utils/parser.py
Extract raw text from PDF and DOCX resume files
"""

import io
import re


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        return f"PDF parsing error: {e}"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        return f"DOCX parsing error: {e}"


def extract_text(uploaded_file) -> str:
    """Route to correct parser based on file type."""
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        return "Unsupported file format. Please upload PDF, DOCX, or TXT."


def extract_email(text: str) -> str:
    match = re.search(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', text)
    return match.group(0) if match else "Not found"


def extract_phone(text: str) -> str:
    match = re.search(r'(\+?\d[\d\s\-().]{7,}\d)', text)
    return match.group(0).strip() if match else "Not found"


def extract_name(text: str) -> str:
    """Heuristic: first non-empty line is usually the candidate name."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        first = lines[0]
        # Likely a name if it's 2-4 words with no special chars
        if len(first.split()) <= 5 and re.match(r'^[A-Za-z .]+$', first):
            return first
    return "Unknown"


def extract_sections(text: str) -> dict:
    """Split resume text into rough sections."""
    section_headers = [
        "summary", "objective", "experience", "education",
        "skills", "projects", "certifications", "achievements",
        "awards", "publications", "languages", "interests"
    ]
    lines = text.split("\n")
    sections = {}
    current_section = "header"
    buffer = []

    for line in lines:
        lowered = line.strip().lower()
        matched = False
        for header in section_headers:
            if lowered.startswith(header) and len(lowered) < 40:
                if buffer:
                    sections[current_section] = "\n".join(buffer)
                current_section = header
                buffer = []
                matched = True
                break
        if not matched:
            buffer.append(line)

    if buffer:
        sections[current_section] = "\n".join(buffer)

    return sections
